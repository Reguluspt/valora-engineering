"""Source-backed, deterministic dossier extraction for Excel and DOCX files."""
from __future__ import annotations

import hashlib
import io
import json
import sys
import tempfile
import unicodedata
import uuid
import zipfile
from dataclasses import dataclass
from datetime import date, datetime, timezone
from pathlib import Path, PurePath
from typing import Any, BinaryIO, Sequence

import docx
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastapi import HTTPException
from sqlalchemy.orm import Session

from app.core.audit import log_audit_event
from app.modules.excel_import.application.adapters import detect_format_and_adapter
from app.modules.excel_import.application.parse_workbook import _is_unsafe_zip_path
from app.modules.excel_import.domain.source_artifact import DEFAULT_SOURCE_LIMITS
from app.modules.excel_import.domain.workbook_adapter import AdapterError, CellValue
from app.modules.excel_import.infrastructure.object_storage import (
    ObjectNotFound,
    ObjectStorageError,
    ObjectStoragePort,
    get_object_storage,
)
from app.modules.excel_import.models import (
    DossierExtractedRow,
    DossierExtractedTable,
    DossierExtractionSnapshot,
    DossierFileRole,
    DossierSourceFile,
    DossierSourceKind,
    DossierTableRole,
    TaskJob,
    TaskJobStatus,
)

_READ_CHUNK = 64 * 1024
_MAX_SOURCE_BYTES = DEFAULT_SOURCE_LIMITS.max_upload_bytes
_MAX_DOCX_ENTRIES = 2_048
_MAX_DOCX_EXPANDED_BYTES = 100 * 1024 * 1024
_MAX_DOCX_TABLES = 500
_MAX_DOCX_ROWS = DEFAULT_SOURCE_LIMITS.max_physical_rows
_MAX_DOCX_COLUMNS = DEFAULT_SOURCE_LIMITS.max_columns
_MAX_CELL_CHARS = DEFAULT_SOURCE_LIMITS.max_cell_chars
_EXTRACTION_SCHEMA_VERSION = "dossier-extraction-v1"
_DOCX_PARSER_NAME = "python-docx"


class DossierExtractionError(Exception):
    def __init__(self, code: str, detail: str, *, retryable: bool = False):
        self.code = code
        self.detail = detail
        self.retryable = retryable
        super().__init__(detail)


@dataclass(frozen=True)
class SourceFingerprint:
    organization_id: uuid.UUID
    dossier_bundle_id: uuid.UUID
    source_file_id: uuid.UUID
    file_role: str
    file_name: str
    file_size_bytes: int
    checksum_sha256: str
    storage_object_key: str

    @classmethod
    def freeze(cls, source: DossierSourceFile) -> SourceFingerprint:
        return cls(
            organization_id=source.organization_id,
            dossier_bundle_id=source.dossier_bundle_id,
            source_file_id=source.id,
            file_role=source.file_role,
            file_name=source.file_name,
            file_size_bytes=source.file_size_bytes,
            checksum_sha256=source.checksum_sha256,
            storage_object_key=source.storage_object_key,
        )


@dataclass(frozen=True)
class ExtractedRowData:
    row_index: int
    is_header: bool
    cells: tuple[dict[str, Any], ...]
    normalized_fields: dict[str, Any]
    locator: dict[str, Any]
    content_fingerprint_sha256: str


@dataclass(frozen=True)
class ExtractedTableData:
    table_index: int
    source_kind: str
    table_role_candidate: str
    role_confidence: float
    raw_title: str | None
    sheet_name: str | None
    page_number: int | None
    header_row_index: int | None
    col_count: int
    locator: dict[str, Any]
    rows: tuple[ExtractedRowData, ...]


@dataclass(frozen=True)
class ExtractedSourceData:
    source_kind: str
    parser_name: str
    parser_version: str
    tables: tuple[ExtractedTableData, ...]
    extraction_digest_sha256: str


def _error(status: int, code: str, detail: str) -> HTTPException:
    return HTTPException(status_code=status, detail={"error_code": code, "detail": detail})


def _canonical_json(value: Any) -> str:
    return json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    )


def _digest(value: Any) -> str:
    return hashlib.sha256(_canonical_json(value).encode("utf-8")).hexdigest()


def _safe_value(value: Any) -> Any:
    if value is None or isinstance(value, (bool, int, str)):
        result = value
    elif isinstance(value, float):
        if value != value or value in {float("inf"), float("-inf")}:
            raise DossierExtractionError("non_finite_number", "Ô dữ liệu chứa số không hữu hạn.")
        result = value
    elif isinstance(value, (date, datetime)):
        result = value.isoformat()
    else:
        result = str(value)
    if isinstance(result, str) and len(result) > _MAX_CELL_CHARS:
        raise DossierExtractionError("cell_length_limit", "Ô dữ liệu vượt quá giới hạn ký tự.")
    return result


def _fold_text(value: Any) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return " ".join(text.casefold().replace("đ", "d").replace("_", " ").split())


_FIELD_ALIASES = {
    "stt": {"stt", "tt", "so thu tu"},
    "name": {"ten tai san", "ten thiet bi", "ten hang muc", "hang muc", "ten"},
    "unit": {"dvt", "don vi tinh", "don vi"},
    "quantity": {"so luong", "sl", "khoi luong"},
    "description": {
        "thong so ky thuat",
        "quy cach",
        "dac tinh",
        "mo ta",
        "model",
        "nhan hieu model",
    },
}


def _normalized_fields(headers: Sequence[Any], values: Sequence[Any]) -> dict[str, Any]:
    mapped: dict[str, Any] = {}
    technical: dict[str, Any] = {}
    for index, raw_header in enumerate(headers):
        header = _fold_text(raw_header)
        value = values[index] if index < len(values) else None
        if value in {None, ""}:
            continue
        target = next(
            (name for name, aliases in _FIELD_ALIASES.items() if header in aliases),
            None,
        )
        if target is not None and target not in mapped:
            mapped[target] = value
        elif header:
            technical[str(raw_header)] = value
    if technical:
        mapped["technical_attributes"] = technical
    return mapped


def _classify_word_table_role(
    title: str | None, headers: Sequence[Any]
) -> tuple[str, float, tuple[str, ...]]:
    sample = _fold_text(" ".join([title or "", *(str(item or "") for item in headers)]))
    rules: list[tuple[DossierTableRole, tuple[str, ...], float]] = [
        (
            DossierTableRole.WORD_TECHNICAL_ASSET_TABLE,
            ("thong so ky thuat", "dac tinh", "quy cach", "cong suat", "model"),
            0.92,
        ),
        (
            DossierTableRole.WORD_QUOTE_COMPARISON_TABLE,
            ("bao gia", "so sanh", "thi truong", "nha cung cap", "don gia"),
            0.90,
        ),
        (
            DossierTableRole.WORD_FINAL_RESULT_TABLE,
            ("ket qua", "gia tri tham dinh", "tong hop", "ket luan", "gia tri con lai"),
            0.88,
        ),
    ]
    candidates: list[tuple[DossierTableRole, list[str], float]] = []
    for role, keywords, confidence in rules:
        matched = [keyword for keyword in keywords if keyword in sample]
        if matched:
            candidates.append((role, matched, min(0.99, confidence + 0.02 * (len(matched) - 1))))
    if not candidates:
        return DossierTableRole.UNKNOWN.value, 0.0, ()
    candidates.sort(key=lambda item: (len(item[1]), item[2]), reverse=True)
    role, matched, confidence = candidates[0]
    return role.value, confidence, tuple(matched)


def _excel_role(headers: Sequence[Any]) -> tuple[str, float]:
    folded = {_fold_text(item) for item in headers if item not in {None, ""}}
    recognized = sum(bool(folded.intersection(aliases)) for aliases in _FIELD_ALIASES.values())
    if recognized >= 2 and (
        folded.intersection(_FIELD_ALIASES["name"])
        or folded.intersection(_FIELD_ALIASES["stt"])
    ):
        return DossierTableRole.EXCEL_CUSTOMER_ASSET_TABLE.value, min(0.99, 0.75 + 0.05 * recognized)
    return DossierTableRole.UNKNOWN.value, 0.20 if folded else 0.0


def _read_verified_source(
    storage: ObjectStoragePort, fingerprint: SourceFingerprint
) -> bytes:
    try:
        stat = storage.head(fingerprint.storage_object_key)
    except ObjectStorageError as exc:
        raise DossierExtractionError(
            "source_head_failed", "Không thể kiểm tra tệp nguồn trong kho lưu trữ.", retryable=True
        ) from exc
    if stat is None:
        raise DossierExtractionError("source_object_missing", "Không tìm thấy tệp nguồn.")
    if stat.size != fingerprint.file_size_bytes:
        raise DossierExtractionError(
            "source_size_mismatch", "Kích thước tệp nguồn không khớp bằng chứng đã lưu."
        )
    if stat.size <= 0 or stat.size > _MAX_SOURCE_BYTES:
        raise DossierExtractionError("source_size_limit", "Kích thước tệp nguồn không hợp lệ.")

    stream: BinaryIO | None = None
    chunks: list[bytes] = []
    total = 0
    try:
        stream = storage.open_stream(fingerprint.storage_object_key)
        while True:
            chunk = stream.read(_READ_CHUNK)
            if not chunk:
                break
            total += len(chunk)
            if total > fingerprint.file_size_bytes or total > _MAX_SOURCE_BYTES:
                raise DossierExtractionError(
                    "source_size_mismatch", "Kích thước tệp nguồn không khớp bằng chứng đã lưu."
                )
            chunks.append(chunk)
    except ObjectNotFound as exc:
        raise DossierExtractionError("source_object_missing", "Không tìm thấy tệp nguồn.") from exc
    except ObjectStorageError as exc:
        raise DossierExtractionError(
            "source_read_failed", "Không thể đọc tệp nguồn từ kho lưu trữ.", retryable=True
        ) from exc
    finally:
        if stream is not None:
            primary_error_active = sys.exc_info()[0] is not None
            try:
                stream.close()
            except Exception as exc:
                if not primary_error_active:
                    raise DossierExtractionError(
                        "source_stream_close_failed",
                        "Không thể đóng luồng tệp nguồn an toàn.",
                        retryable=True,
                    ) from exc
    data = b"".join(chunks)
    if len(data) != fingerprint.file_size_bytes:
        raise DossierExtractionError(
            "source_size_mismatch", "Kích thước tệp nguồn không khớp bằng chứng đã lưu."
        )
    if hashlib.sha256(data).hexdigest() != fingerprint.checksum_sha256:
        raise DossierExtractionError(
            "source_checksum_mismatch", "Mã kiểm tra tệp nguồn không khớp bằng chứng đã lưu."
        )
    return data


def _validate_docx_archive(data: bytes) -> None:
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except (zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        raise DossierExtractionError("invalid_docx", "Tệp DOCX không hợp lệ.") from exc
    with archive:
        entries = archive.infolist()
        if len(entries) > _MAX_DOCX_ENTRIES:
            raise DossierExtractionError("docx_entry_limit", "Tệp DOCX có quá nhiều thành phần.")
        if sum(item.file_size for item in entries) > _MAX_DOCX_EXPANDED_BYTES:
            raise DossierExtractionError(
                "docx_expansion_limit", "Kích thước giải nén DOCX vượt giới hạn."
            )
        names = {item.filename for item in entries}
        if not {"[Content_Types].xml", "word/document.xml"} <= names:
            raise DossierExtractionError("invalid_docx", "Tệp DOCX thiếu cấu trúc bắt buộc.")
        for item in entries:
            if _is_unsafe_zip_path(item.filename):
                raise DossierExtractionError(
                    "unsafe_docx_path", "Tệp DOCX chứa đường dẫn không an toàn."
                )
            if item.flag_bits & 0x1:
                raise DossierExtractionError(
                    "encrypted_docx", "Tệp DOCX mã hóa không được hỗ trợ."
                )
            lowered = item.filename.casefold().replace("\\", "/")
            if "vbaproject" in lowered or lowered.endswith(".bin"):
                raise DossierExtractionError(
                    "docx_active_content_forbidden", "Tệp DOCX chứa nội dung chủ động."
                )


def _docx_page_breaks(paragraph: Paragraph) -> int:
    return len(paragraph._p.xpath(".//w:br[@w:type='page']")) + len(
        paragraph._p.xpath(".//w:lastRenderedPageBreak")
    )


def _build_row(
    *,
    row_index: int,
    is_header: bool,
    values: Sequence[Any],
    headers: Sequence[Any],
    cells: Sequence[dict[str, Any]],
    locator: dict[str, Any],
) -> ExtractedRowData:
    normalized = {} if is_header else _normalized_fields(headers, values)
    fingerprint = _digest({"values": list(values), "locator": locator})
    return ExtractedRowData(
        row_index=row_index,
        is_header=is_header,
        cells=tuple(cells),
        normalized_fields=normalized,
        locator=locator,
        content_fingerprint_sha256=fingerprint,
    )


def parse_docx_bytes(data: bytes) -> ExtractedSourceData:
    _validate_docx_archive(data)
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise DossierExtractionError("invalid_docx", "Không thể đọc nội dung DOCX.") from exc

    tables: list[ExtractedTableData] = []
    last_title: str | None = None
    page_number = 1
    section_index = 0
    for block in document.iter_inner_content():
        if isinstance(block, Paragraph):
            text = " ".join(block.text.split())
            if text:
                last_title = text[:255]
            page_number += _docx_page_breaks(block)
            if block._p.xpath("./w:pPr/w:sectPr"):
                section_index += 1
            continue
        if not isinstance(block, Table):
            continue
        if len(tables) >= _MAX_DOCX_TABLES:
            raise DossierExtractionError("docx_table_limit", "DOCX có quá nhiều bảng.")
        table_index = len(tables)
        raw_values: list[list[Any]] = []
        raw_cells: list[list[dict[str, Any]]] = []
        max_columns = 0
        for row_index, row in enumerate(block.rows):
            if row_index >= _MAX_DOCX_ROWS:
                raise DossierExtractionError("docx_row_limit", "Bảng DOCX có quá nhiều dòng.")
            if len(row.cells) > _MAX_DOCX_COLUMNS:
                raise DossierExtractionError("docx_column_limit", "Bảng DOCX có quá nhiều cột.")
            values: list[Any] = []
            cells: list[dict[str, Any]] = []
            for cell_index, cell in enumerate(row.cells):
                value = _safe_value("\n".join(part.text for part in cell.paragraphs).strip())
                values.append(value)
                cells.append(
                    {
                        "cell_index": cell_index,
                        "value": value,
                        "cell_type": "string",
                        "locator": {
                            "part": "word/document.xml",
                            "table_index": table_index,
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "page_number": page_number,
                        },
                    }
                )
            max_columns = max(max_columns, len(values))
            raw_values.append(values)
            raw_cells.append(cells)
        headers = raw_values[0] if raw_values else []
        role, confidence, matched_keywords = _classify_word_table_role(last_title, headers)
        rows = tuple(
            _build_row(
                row_index=row_index,
                is_header=row_index == 0,
                values=values,
                headers=headers,
                cells=raw_cells[row_index],
                locator={
                    "part": "word/document.xml",
                    "table_index": table_index,
                    "row_index": row_index,
                    "page_number": page_number,
                    "page_locator_kind": "rendered_break_estimate",
                },
            )
            for row_index, values in enumerate(raw_values)
        )
        tables.append(
            ExtractedTableData(
                table_index=table_index,
                source_kind=DossierSourceKind.DOCX.value,
                table_role_candidate=role,
                role_confidence=confidence,
                raw_title=last_title,
                sheet_name=None,
                page_number=page_number,
                header_row_index=0 if rows else None,
                col_count=max_columns,
                locator={
                    "part": "word/document.xml",
                    "table_index": table_index,
                    "section_index": section_index,
                    "page_number": page_number,
                    "page_locator_kind": "rendered_break_estimate",
                    "role_rule_keywords": list(matched_keywords),
                },
                rows=rows,
            )
        )
        last_title = None
    return _finalize_source(
        source_kind=DossierSourceKind.DOCX.value,
        parser_name=_DOCX_PARSER_NAME,
        parser_version=docx.__version__,
        tables=tables,
    )


def _excel_cell(cell: CellValue, *, sheet_name: str) -> dict[str, Any]:
    return {
        "cell_index": cell.column - 1,
        "value": _safe_value(cell.value),
        "cell_type": cell.cell_type,
        "locator": {
            "sheet_name": sheet_name,
            "coordinate": cell.coordinate,
            "row_number": cell.row,
            "column_number": cell.column,
        },
    }


def _nonempty(cells: Sequence[CellValue]) -> bool:
    return any(cell.value not in {None, ""} for cell in cells)


def _excel_blocks(rows: Sequence[Sequence[CellValue]]) -> list[list[Sequence[CellValue]]]:
    blocks: list[list[Sequence[CellValue]]] = []
    current: list[Sequence[CellValue]] = []
    for row in rows:
        if _nonempty(row):
            current.append(row)
        elif current:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)
    return blocks


def parse_excel_bytes(data: bytes, *, filename: str) -> ExtractedSourceData:
    suffix = PurePath(filename).suffix.casefold()
    with tempfile.TemporaryDirectory(prefix="valora-dossier-") as temp_dir:
        path = Path(temp_dir) / f"source{suffix}"
        with path.open("wb") as output:
            output.write(data)
        adapter = None
        try:
            detected_format, adapter = detect_format_and_adapter(
                str(path), filename, limits=DEFAULT_SOURCE_LIMITS
            )
            inspection = adapter.inspect(str(path))
            tables: list[ExtractedTableData] = []
            for sheet_name in inspection.sheet_names:
                sheet_rows = list(adapter.iter_rows(str(path), sheet_name))
                for block in _excel_blocks(sheet_rows):
                    table_index = len(tables)
                    headers = [_safe_value(cell.value) for cell in block[0]] if block else []
                    role, confidence = _excel_role(headers)
                    extracted_rows: list[ExtractedRowData] = []
                    for row_index, source_row in enumerate(block):
                        values = [_safe_value(cell.value) for cell in source_row]
                        cells = [_excel_cell(cell, sheet_name=sheet_name) for cell in source_row]
                        extracted_rows.append(
                            _build_row(
                                row_index=row_index,
                                is_header=row_index == 0,
                                values=values,
                                headers=headers,
                                cells=cells,
                                locator={
                                    "sheet_name": sheet_name,
                                    "source_row_number": source_row[0].row if source_row else None,
                                    "start_coordinate": (
                                        source_row[0].coordinate if source_row else None
                                    ),
                                    "end_coordinate": (
                                        source_row[-1].coordinate if source_row else None
                                    ),
                                },
                            )
                        )
                    first_row = block[0][0].row if block and block[0] else None
                    last_row = block[-1][0].row if block and block[-1] else None
                    tables.append(
                        ExtractedTableData(
                            table_index=table_index,
                            source_kind=DossierSourceKind.EXCEL.value,
                            table_role_candidate=role,
                            role_confidence=confidence,
                            raw_title=sheet_name[:255],
                            sheet_name=sheet_name[:255],
                            page_number=None,
                            header_row_index=0 if extracted_rows else None,
                            col_count=max((len(row) for row in block), default=0),
                            locator={
                                "sheet_name": sheet_name,
                                "start_row": first_row,
                                "end_row": last_row,
                                "table_index": table_index,
                            },
                            rows=tuple(extracted_rows),
                        )
                    )
            return _finalize_source(
                source_kind=DossierSourceKind.EXCEL.value,
                parser_name=inspection.adapter_name,
                parser_version=inspection.adapter_version,
                tables=tables,
            )
        except AdapterError as exc:
            raise DossierExtractionError(exc.error_code, exc.detail) from exc
        finally:
            if adapter is not None:
                adapter.close()


def _finalize_source(
    *,
    source_kind: str,
    parser_name: str,
    parser_version: str,
    tables: Sequence[ExtractedTableData],
) -> ExtractedSourceData:
    serializable = [
        {
            "table_index": table.table_index,
            "source_kind": table.source_kind,
            "role": table.table_role_candidate,
            "confidence": table.role_confidence,
            "raw_title": table.raw_title,
            "sheet_name": table.sheet_name,
            "page_number": table.page_number,
            "header_row_index": table.header_row_index,
            "col_count": table.col_count,
            "locator": table.locator,
            "rows": [
                {
                    "row_index": row.row_index,
                    "is_header": row.is_header,
                    "cells": list(row.cells),
                    "normalized_fields": row.normalized_fields,
                    "locator": row.locator,
                    "content_fingerprint_sha256": row.content_fingerprint_sha256,
                }
                for row in table.rows
            ],
        }
        for table in tables
    ]
    return ExtractedSourceData(
        source_kind=source_kind,
        parser_name=parser_name,
        parser_version=parser_version,
        tables=tuple(tables),
        extraction_digest_sha256=_digest(
            {"schema_version": _EXTRACTION_SCHEMA_VERSION, "tables": serializable}
        ),
    )


def _source_unchanged(source: DossierSourceFile, frozen: SourceFingerprint) -> bool:
    return SourceFingerprint.freeze(source) == frozen


def _assert_live_job(
    job: TaskJob | None,
    *,
    dossier_bundle_id: uuid.UUID,
    source_file_id: uuid.UUID,
    generation_token: int,
) -> None:
    now = datetime.now(timezone.utc)
    lease = job.lease_expires_at if job is not None else None
    if lease is not None and lease.tzinfo is None:
        lease = lease.replace(tzinfo=timezone.utc)
    payload = job.payload if job is not None else {}
    if (
        job is None
        or job.job_type != "document_extraction"
        or job.status != TaskJobStatus.CLAIMED.value
        or job.generation_token != generation_token
        or lease is None
        or lease <= now
        or payload.get("dossier_bundle_id") != str(dossier_bundle_id)
        or payload.get("source_file_id") != str(source_file_id)
    ):
        raise DossierExtractionError(
            "stale_extraction_job", "Job extraction không còn lease/generation hợp lệ."
        )


def extract_dossier_source(
    db: Session,
    *,
    org_id: uuid.UUID,
    dossier_bundle_id: uuid.UUID,
    source_file_id: uuid.UUID,
    job_id: uuid.UUID,
    generation_token: int,
    storage: ObjectStoragePort | None = None,
) -> DossierExtractionSnapshot:
    """Read verified object bytes, parse real content and append one immutable snapshot."""
    source = (
        db.query(DossierSourceFile)
        .filter(
            DossierSourceFile.organization_id == org_id,
            DossierSourceFile.dossier_bundle_id == dossier_bundle_id,
            DossierSourceFile.id == source_file_id,
        )
        .populate_existing()
        .first()
    )
    if source is None:
        raise _error(404, "resource_not_found", "Không tìm thấy tệp nguồn.")
    frozen = SourceFingerprint.freeze(source)
    suffix = PurePath(frozen.file_name).suffix.casefold()
    if frozen.file_role == DossierFileRole.CUSTOMER_ASSET_LIST.value and suffix in {
        ".xls",
        ".xlsx",
    }:
        source_kind = DossierSourceKind.EXCEL.value
    elif (
        frozen.file_role == DossierFileRole.FINAL_APPRAISAL_REPORT.value and suffix == ".docx"
    ):
        source_kind = DossierSourceKind.DOCX.value
    elif frozen.file_role == DossierFileRole.FINAL_APPRAISAL_REPORT.value and suffix == ".pdf":
        raise DossierExtractionError(
            "pdf_extraction_not_implemented",
            "PDF cần runtime layout/OCR riêng; không được thay bằng dữ liệu mô phỏng.",
        )
    else:
        raise DossierExtractionError(
            "unsupported_source_role_format", "Vai trò và định dạng tệp không hỗ trợ extraction."
        )

    db.rollback()
    raw_bytes = _read_verified_source(storage or get_object_storage(), frozen)
    extracted = (
        parse_excel_bytes(raw_bytes, filename=frozen.file_name)
        if source_kind == DossierSourceKind.EXCEL.value
        else parse_docx_bytes(raw_bytes)
    )

    job = (
        db.query(TaskJob)
        .filter(TaskJob.organization_id == org_id, TaskJob.id == job_id)
        .populate_existing()
        .with_for_update()
        .first()
    )
    _assert_live_job(
        job,
        dossier_bundle_id=dossier_bundle_id,
        source_file_id=source_file_id,
        generation_token=generation_token,
    )
    locked_source = (
        db.query(DossierSourceFile)
        .filter(
            DossierSourceFile.organization_id == org_id,
            DossierSourceFile.dossier_bundle_id == dossier_bundle_id,
            DossierSourceFile.id == source_file_id,
        )
        .populate_existing()
        .with_for_update()
        .first()
    )
    if locked_source is None or not _source_unchanged(locked_source, frozen):
        raise DossierExtractionError(
            "source_changed_during_extraction", "Bằng chứng tệp nguồn đã thay đổi khi extraction."
        )
    existing = (
        db.query(DossierExtractionSnapshot)
        .filter(
            DossierExtractionSnapshot.organization_id == org_id,
            DossierExtractionSnapshot.dossier_bundle_id == dossier_bundle_id,
            DossierExtractionSnapshot.source_file_id == source_file_id,
            DossierExtractionSnapshot.source_checksum_sha256 == frozen.checksum_sha256,
            DossierExtractionSnapshot.parser_name == extracted.parser_name,
            DossierExtractionSnapshot.parser_version == extracted.parser_version,
            DossierExtractionSnapshot.extraction_schema_version == _EXTRACTION_SCHEMA_VERSION,
        )
        .first()
    )
    if existing is not None:
        db.commit()
        return existing

    snapshot = DossierExtractionSnapshot(
        organization_id=org_id,
        dossier_bundle_id=dossier_bundle_id,
        source_file_id=source_file_id,
        source_checksum_sha256=frozen.checksum_sha256,
        source_kind=extracted.source_kind,
        parser_name=extracted.parser_name,
        parser_version=extracted.parser_version,
        extraction_schema_version=_EXTRACTION_SCHEMA_VERSION,
        created_by_job_id=job_id,
        generation_token=generation_token,
        table_count=len(extracted.tables),
        row_count=sum(len(table.rows) for table in extracted.tables),
        extraction_digest_sha256=extracted.extraction_digest_sha256,
    )
    db.add(snapshot)
    db.flush()
    for table_data in extracted.tables:
        table = DossierExtractedTable(
            organization_id=org_id,
            dossier_bundle_id=dossier_bundle_id,
            source_file_id=source_file_id,
            extraction_snapshot_id=snapshot.id,
            table_index=table_data.table_index,
            source_kind=table_data.source_kind,
            table_role_candidate=table_data.table_role_candidate,
            role_confidence=table_data.role_confidence,
            raw_title=table_data.raw_title,
            sheet_name=table_data.sheet_name,
            page_number=table_data.page_number,
            header_row_index=table_data.header_row_index,
            row_count=len(table_data.rows),
            col_count=table_data.col_count,
            locator_json=table_data.locator,
        )
        db.add(table)
        db.flush()
        db.add_all(
            [
                DossierExtractedRow(
                    organization_id=org_id,
                    dossier_bundle_id=dossier_bundle_id,
                    source_file_id=source_file_id,
                    extracted_table_id=table.id,
                    row_index=row.row_index,
                    is_header=row.is_header,
                    cells_json=list(row.cells),
                    normalized_fields=row.normalized_fields,
                    locator_json=row.locator,
                    content_fingerprint_sha256=row.content_fingerprint_sha256,
                )
                for row in table_data.rows
            ]
        )
    log_audit_event(
        db,
        event_name="DossierSourceExtracted",
        entity_type="DossierExtractionSnapshot",
        entity_id=snapshot.id,
        organization_id=org_id,
        command_name="ExtractDossierDocument",
        correlation_id=job.correlation_id,
        payload={
            "dossier_bundle_id": str(dossier_bundle_id),
            "source_file_id": str(source_file_id),
            "source_checksum_sha256": frozen.checksum_sha256,
            "source_kind": extracted.source_kind,
            "parser_name": extracted.parser_name,
            "parser_version": extracted.parser_version,
            "extraction_schema_version": _EXTRACTION_SCHEMA_VERSION,
            "table_count": snapshot.table_count,
            "row_count": snapshot.row_count,
            "extraction_digest_sha256": snapshot.extraction_digest_sha256,
            "generation_token": generation_token,
        },
    )
    db.commit()
    db.refresh(snapshot)
    return snapshot
