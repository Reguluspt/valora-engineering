"""S15-R-002 tests using real XLSX/DOCX bytes and source locators."""
from __future__ import annotations

import hashlib
import io
from dataclasses import dataclass

import pytest
from docx import Document
from openpyxl import Workbook
from sqlalchemy import create_engine, event
from sqlalchemy.orm import Session
from sqlalchemy.pool import StaticPool

from app.db import Base
from app.modules.document_engine_intelligence.application.dossier_bundle_service import (
    DossierSourceSpec,
    create_paired_dossier_bundle,
)
from app.modules.document_engine_intelligence.application.dossier_extraction_service import (
    DossierExtractionError,
    extract_dossier_source,
)
from app.modules.excel_import.infrastructure.object_storage import FakeObjectStorage
from app.modules.excel_import.models import (
    DossierExtractedRow,
    DossierExtractedTable,
    DossierExtractionSnapshot,
    DossierSourceFile,
    DossierTableRole,
)
from app.modules.project_master_data.models import (
    AuditEvent,
    Customer,
    OrganizationProfile,
    OrganizationStatus,
    Role,
    User,
    UserRole,
    UserStatus,
)
from app.modules.workflow_workbench.application.reliable_job_service import (
    claim_job_lease,
    complete_job,
    enqueue_durable_job,
)


@dataclass(frozen=True)
class SourceBackedSeed:
    actor: User
    bundle_id: object
    excel_source: DossierSourceFile
    report_source: DossierSourceFile
    storage: FakeObjectStorage
    excel_bytes: bytes
    docx_bytes: bytes


@pytest.fixture
def db_session() -> Session:
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    @event.listens_for(engine, "connect")
    def _enable_foreign_keys(dbapi_connection, _connection_record) -> None:
        cursor = dbapi_connection.cursor()
        cursor.execute("PRAGMA foreign_keys=ON")
        cursor.close()

    Base.metadata.create_all(engine)
    session = Session(engine)
    try:
        yield session
    finally:
        session.close()
        engine.dispose()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Danh sách tài sản"
    sheet.append(["STT", "Tên tài sản", "ĐVT", "Số lượng", "Thông số kỹ thuật"])
    sheet.append([1, "Máy biến áp ABB 110kV", "cái", 2, "63 MVA"])
    sheet.append([2, "Máy phát điện Cummins C500", "cái", 1, "500 kVA"])
    sheet.append([])
    sheet.append(["Ghi chú hiện trường chưa có tiêu đề chuẩn"])
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _add_table(document: Document, title: str, rows: list[list[object]]) -> None:
    document.add_paragraph(title)
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = str(value)


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("BÁO CÁO THẨM ĐỊNH GIÁ", level=1)
    _add_table(
        document,
        "Bảng thông số kỹ thuật tài sản",
        [
            ["STT", "Tên thiết bị", "ĐVT", "Số lượng", "Thông số kỹ thuật"],
            [2, "Máy phát điện Cummins C500", "bộ", 1, "500 kVA"],
            [1, "Máy biến áp ABB 110kV", "cái", 2, "63 MVA"],
        ],
    )
    _add_table(
        document,
        "Bảng so sánh báo giá thị trường",
        [
            ["STT", "Tên thiết bị", "Nhà cung cấp", "Đơn giá"],
            [1, "Máy biến áp ABB 110kV", "Nhà cung cấp A", "1250000000"],
            [2, "Máy phát điện Cummins C500", "Nhà cung cấp B", "450000000"],
        ],
    )
    _add_table(
        document,
        "Bảng tổng hợp kết quả định giá",
        [
            ["STT", "Tên tài sản", "ĐVT", "Số lượng", "Giá trị thẩm định"],
            [1, "Máy biến áp ABB 110kV", "cái", 2, "2400000000"],
            [2, "Máy phát điện Cummins C500", "cái", 1, "420000000"],
        ],
    )
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _seed_source_backed_dossier(db: Session, slug: str = "real") -> SourceBackedSeed:
    organization = OrganizationProfile(
        legal_name=f"Organization {slug}",
        organization_slug=slug,
        status=OrganizationStatus.ACTIVE,
    )
    role = Role(
        code=f"extractor-{slug}",
        display_name=f"Extractor {slug}",
        permissions=["document_intelligence:job:create"],
    )
    db.add_all([organization, role])
    db.commit()
    actor = User(
        organization_id=organization.id,
        email=f"extractor-{slug}@example.test",
        full_name=f"Extractor {slug}",
        status=UserStatus.ACTIVE,
    )
    db.add(actor)
    db.commit()
    db.add(UserRole(user_id=actor.id, role_id=role.id, is_active=True))
    db.commit()
    customer = Customer(
        organization_id=organization.id,
        legal_name=f"Customer {slug}",
        status="active",
        created_by=actor.id,
    )
    db.add(customer)
    db.commit()

    excel_bytes = _xlsx_bytes()
    docx_bytes = _docx_bytes()
    excel_key = f"verified/{slug}/assets.xlsx"
    report_key = f"verified/{slug}/report.docx"
    storage = FakeObjectStorage()
    storage.put_stream(
        excel_key,
        io.BytesIO(excel_bytes),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        expected_size=len(excel_bytes),
    )
    storage.put_stream(
        report_key,
        io.BytesIO(docx_bytes),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        expected_size=len(docx_bytes),
    )
    bundle, sources = create_paired_dossier_bundle(
        db,
        actor=actor,
        org_id=organization.id,
        customer_id=customer.id,
        project_id=None,
        bundle_code=f"HS-{slug}",
        files=[
            DossierSourceSpec(
                file_role="customer_asset_list",
                file_name="assets.xlsx",
                file_size_bytes=len(excel_bytes),
                checksum_sha256=hashlib.sha256(excel_bytes).hexdigest(),
                storage_object_key=excel_key,
            ),
            DossierSourceSpec(
                file_role="final_appraisal_report",
                file_name="report.docx",
                file_size_bytes=len(docx_bytes),
                checksum_sha256=hashlib.sha256(docx_bytes).hexdigest(),
                storage_object_key=report_key,
            ),
        ],
    )
    by_role = {source.file_role: source for source in sources}
    return SourceBackedSeed(
        actor=actor,
        bundle_id=bundle.id,
        excel_source=by_role["customer_asset_list"],
        report_source=by_role["final_appraisal_report"],
        storage=storage,
        excel_bytes=excel_bytes,
        docx_bytes=docx_bytes,
    )


def _claim_extraction_job(
    db: Session,
    seed: SourceBackedSeed,
    source: DossierSourceFile,
    *,
    key: str,
) -> tuple[object, object]:
    job = enqueue_durable_job(
        db,
        actor=seed.actor,
        org_id=seed.actor.organization_id,
        job_type="document_extraction",
        idempotency_key=key,
        payload={
            "dossier_bundle_id": str(seed.bundle_id),
            "source_file_id": str(source.id),
        },
    )
    db.commit()
    return claim_job_lease(
        db,
        worker_id="source-worker",
        job_id=job.id,
        org_id=seed.actor.organization_id,
        lease_duration_seconds=60,
    )


def _extract_and_complete(
    db: Session,
    seed: SourceBackedSeed,
    source: DossierSourceFile,
    *,
    key: str,
) -> DossierExtractionSnapshot:
    job, attempt = _claim_extraction_job(db, seed, source, key=key)
    snapshot = extract_dossier_source(
        db,
        org_id=seed.actor.organization_id,
        dossier_bundle_id=seed.bundle_id,
        source_file_id=source.id,
        job_id=job.id,
        generation_token=job.generation_token,
        storage=seed.storage,
    )
    complete_job(
        db,
        worker_id="source-worker",
        job_id=job.id,
        org_id=seed.actor.organization_id,
        attempt_id=attempt.id,
        generation_token=job.generation_token,
        result_payload={"extraction_snapshot_id": str(snapshot.id)},
    )
    return snapshot


def test_real_xlsx_and_docx_are_extracted_with_locators(db_session: Session) -> None:
    seed = _seed_source_backed_dossier(db_session, "real-content")
    excel_snapshot = _extract_and_complete(
        db_session, seed, seed.excel_source, key="extract-real-excel"
    )
    report_snapshot = _extract_and_complete(
        db_session, seed, seed.report_source, key="extract-real-docx"
    )

    assert excel_snapshot.source_kind == "excel"
    assert excel_snapshot.table_count == 2
    assert report_snapshot.source_kind == "docx"
    assert report_snapshot.table_count == 3
    roles = {
        table.table_role_candidate
        for table in db_session.query(DossierExtractedTable)
        .filter_by(extraction_snapshot_id=report_snapshot.id)
        .all()
    }
    assert roles == {
        DossierTableRole.WORD_TECHNICAL_ASSET_TABLE.value,
        DossierTableRole.WORD_QUOTE_COMPARISON_TABLE.value,
        DossierTableRole.WORD_FINAL_RESULT_TABLE.value,
    }
    excel_rows = (
        db_session.query(DossierExtractedRow)
        .join(DossierExtractedTable)
        .filter(
            DossierExtractedTable.extraction_snapshot_id == excel_snapshot.id,
            DossierExtractedRow.is_header.is_(False),
        )
        .all()
    )
    excel_row = next(
        (row for row in excel_rows if str(row.normalized_fields.get("stt")) == "1"),
        None,
    )
    assert excel_row is not None
    assert excel_row.normalized_fields["name"] == "Máy biến áp ABB 110kV"
    assert excel_row.locator_json["sheet_name"] == "Danh sách tài sản"
    assert excel_row.cells_json[0]["locator"]["coordinate"] == "A2"
    report_row = (
        db_session.query(DossierExtractedRow)
        .join(DossierExtractedTable)
        .filter(
            DossierExtractedTable.extraction_snapshot_id == report_snapshot.id,
            DossierExtractedRow.is_header.is_(False),
        )
        .first()
    )
    assert report_row is not None
    assert report_row.locator_json["part"] == "word/document.xml"
    assert "table_index" in report_row.cells_json[0]["locator"]
    assert (
        db_session.query(AuditEvent)
        .filter_by(event_name="DossierSourceExtracted")
        .count()
        == 2
    )


def test_extraction_rejects_checksum_mismatch_without_persisting_rows(
    db_session: Session,
) -> None:
    seed = _seed_source_backed_dossier(db_session, "checksum")
    job, _attempt = _claim_extraction_job(
        db_session, seed, seed.report_source, key="extract-bad-checksum"
    )
    wrong_storage = FakeObjectStorage()
    wrong_storage.put_stream(
        seed.report_source.storage_object_key,
        io.BytesIO(b"x" * len(seed.docx_bytes)),
        content_type="application/octet-stream",
        expected_size=len(seed.docx_bytes),
    )
    with pytest.raises(DossierExtractionError) as exc_info:
        extract_dossier_source(
            db_session,
            org_id=seed.actor.organization_id,
            dossier_bundle_id=seed.bundle_id,
            source_file_id=seed.report_source.id,
            job_id=job.id,
            generation_token=job.generation_token,
            storage=wrong_storage,
        )
    assert exc_info.value.code == "source_checksum_mismatch"
    db_session.rollback()
    assert db_session.query(DossierExtractionSnapshot).count() == 0
    assert db_session.query(DossierExtractedRow).count() == 0


def test_extraction_rerun_reuses_exact_snapshot(db_session: Session) -> None:
    seed = _seed_source_backed_dossier(db_session, "idempotent-extraction")
    first = _extract_and_complete(
        db_session, seed, seed.report_source, key="extract-idempotent-1"
    )
    second = _extract_and_complete(
        db_session, seed, seed.report_source, key="extract-idempotent-2"
    )
    assert second.id == first.id
    assert db_session.query(DossierExtractionSnapshot).count() == 1
    assert db_session.query(AuditEvent).filter_by(event_name="DossierSourceExtracted").count() == 1
