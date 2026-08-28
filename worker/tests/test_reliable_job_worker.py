"""End-to-end SQLite tests for the reliable job consumer."""
from __future__ import annotations

from collections.abc import Iterator
from dataclasses import dataclass
import hashlib
import io

from docx import Document
from openpyxl import Workbook
import pytest
from sqlalchemy import create_engine, event
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.db import Base
from app.modules.document_engine_intelligence.application.dossier_bundle_service import (
    DossierSourceSpec,
    create_paired_dossier_bundle,
)
from app.modules.excel_import.infrastructure.object_storage import FakeObjectStorage
from app.modules.excel_import.models import (
    DossierAlignmentRun,
    DossierExtractionSnapshot,
    DossierRowAlignment,
    TaskJob,
    TaskJobAttempt,
)
from app.modules.project_master_data.models import (
    AuditEvent,
    Customer,
    OrganizationProfile,
    OrganizationStatus,
    Role,
    User,
    UserRole,
    UserStatus,
)
from app.modules.workflow_workbench.application.reliable_job_service import (
    enqueue_durable_job,
)
from worker.handlers import build_handler_registry
from worker.runtime import JobExecutionContext, JobHandlerFailure, ReliableJobWorker


@dataclass(frozen=True)
class RuntimeSeed:
    session_factory: sessionmaker[Session]
    job_id: object


@dataclass(frozen=True)
class SourceBackedRuntimeSeed:
    organization_id: object
    actor_id: object
    bundle_id: object
    storage: FakeObjectStorage


@pytest.fixture
def session_factory() -> Iterator[sessionmaker[Session]]:
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    @event.listens_for(engine, "connect")
    def _enable_foreign_keys(dbapi_connection, _connection_record) -> None:
        cursor = dbapi_connection.cursor()
        cursor.execute("PRAGMA foreign_keys=ON")
        cursor.close()

    Base.metadata.create_all(engine)
    factory = sessionmaker(bind=engine, expire_on_commit=False)
    try:
        yield factory
    finally:
        engine.dispose()


def _seed_job(
    factory: sessionmaker[Session], *, slug: str, max_attempts: int = 3
) -> RuntimeSeed:
    with factory() as db:
        organization = OrganizationProfile(
            legal_name=f"Organization {slug}",
            organization_slug=slug,
            status=OrganizationStatus.ACTIVE,
        )
        role = Role(
            code=f"worker-role-{slug}",
            display_name=f"Worker role {slug}",
            permissions=["document_intelligence:job:create"],
        )
        db.add_all([organization, role])
        db.commit()
        actor = User(
            organization_id=organization.id,
            email=f"actor-{slug}@example.test",
            full_name=f"Actor {slug}",
            status=UserStatus.ACTIVE,
        )
        db.add(actor)
        db.commit()
        db.add(UserRole(user_id=actor.id, role_id=role.id, is_active=True))
        db.commit()
        customer = Customer(
            organization_id=organization.id,
            legal_name=f"Customer {slug}",
            status="active",
            created_by=actor.id,
        )
        db.add(customer)
        db.commit()
        bundle, sources = create_paired_dossier_bundle(
            db,
            actor=actor,
            org_id=organization.id,
            customer_id=customer.id,
            project_id=None,
            bundle_code=f"HS-{slug}",
            files=[
                DossierSourceSpec(
                    file_role="customer_asset_list",
                    file_name="assets.xlsx",
                    file_size_bytes=100,
                    checksum_sha256="a" * 64,
                    storage_object_key=f"verified/{slug}/assets.xlsx",
                ),
                DossierSourceSpec(
                    file_role="final_appraisal_report",
                    file_name="report.docx",
                    file_size_bytes=200,
                    checksum_sha256="b" * 64,
                    storage_object_key=f"verified/{slug}/report.docx",
                ),
            ],
        )
        job = enqueue_durable_job(
            db,
            actor=actor,
            org_id=organization.id,
            job_type="document_extraction",
            idempotency_key=f"worker-job-{slug}",
            payload={
                "dossier_bundle_id": str(bundle.id),
                "source_file_id": str(sources[1].id),
            },
            max_attempts=max_attempts,
        )
        db.commit()
        return RuntimeSeed(factory, job.id)


def _real_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Danh sách tài sản"
    sheet.append(["STT", "Tên tài sản", "ĐVT", "Số lượng", "Thông số kỹ thuật"])
    sheet.append([1, "Máy biến áp ABB 110kV", "cái", 2, "63 MVA"])
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _real_docx_bytes() -> bytes:
    document = Document()
    for title, rows in [
        (
            "Bảng thông số kỹ thuật tài sản",
            [
                ["STT", "Tên thiết bị", "ĐVT", "Số lượng", "Thông số kỹ thuật"],
                [1, "Máy biến áp ABB 110kV", "cái", 2, "63 MVA"],
            ],
        ),
        (
            "Bảng so sánh báo giá thị trường",
            [
                ["STT", "Tên thiết bị", "Nhà cung cấp", "Đơn giá"],
                [1, "Máy biến áp ABB 110kV", "Nhà cung cấp A", 1_250_000_000],
            ],
        ),
        (
            "Bảng tổng hợp kết quả định giá",
            [
                ["STT", "Tên tài sản", "ĐVT", "Số lượng", "Giá trị thẩm định"],
                [1, "Máy biến áp ABB 110kV", "cái", 2, 2_400_000_000],
            ],
        ),
    ]:
        document.add_paragraph(title)
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for row_index, row in enumerate(rows):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = str(value)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _seed_source_backed_jobs(
    factory: sessionmaker[Session], *, slug: str
) -> SourceBackedRuntimeSeed:
    excel_bytes = _real_xlsx_bytes()
    docx_bytes = _real_docx_bytes()
    storage = FakeObjectStorage()
    excel_key = f"verified/{slug}/assets.xlsx"
    report_key = f"verified/{slug}/report.docx"
    storage.put_stream(
        excel_key,
        io.BytesIO(excel_bytes),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        expected_size=len(excel_bytes),
    )
    storage.put_stream(
        report_key,
        io.BytesIO(docx_bytes),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        expected_size=len(docx_bytes),
    )
    with factory() as db:
        organization = OrganizationProfile(
            legal_name=f"Organization {slug}",
            organization_slug=slug,
            status=OrganizationStatus.ACTIVE,
        )
        role = Role(
            code=f"worker-role-{slug}",
            display_name=f"Worker role {slug}",
            permissions=["document_intelligence:job:create"],
        )
        db.add_all([organization, role])
        db.commit()
        actor = User(
            organization_id=organization.id,
            email=f"actor-{slug}@example.test",
            full_name=f"Actor {slug}",
            status=UserStatus.ACTIVE,
        )
        db.add(actor)
        db.commit()
        db.add(UserRole(user_id=actor.id, role_id=role.id, is_active=True))
        db.commit()
        customer = Customer(
            organization_id=organization.id,
            legal_name=f"Customer {slug}",
            status="active",
            created_by=actor.id,
        )
        db.add(customer)
        db.commit()
        bundle, sources = create_paired_dossier_bundle(
            db,
            actor=actor,
            org_id=organization.id,
            customer_id=customer.id,
            project_id=None,
            bundle_code=f"HS-{slug}",
            files=[
                DossierSourceSpec(
                    file_role="customer_asset_list",
                    file_name="assets.xlsx",
                    file_size_bytes=len(excel_bytes),
                    checksum_sha256=hashlib.sha256(excel_bytes).hexdigest(),
                    storage_object_key=excel_key,
                ),
                DossierSourceSpec(
                    file_role="final_appraisal_report",
                    file_name="report.docx",
                    file_size_bytes=len(docx_bytes),
                    checksum_sha256=hashlib.sha256(docx_bytes).hexdigest(),
                    storage_object_key=report_key,
                ),
            ],
        )
        for index, source in enumerate(sources, start=1):
            enqueue_durable_job(
                db,
                actor=actor,
                org_id=organization.id,
                job_type="document_extraction",
                idempotency_key=f"source-backed-{slug}-{index}",
                payload={
                    "dossier_bundle_id": str(bundle.id),
                    "source_file_id": str(source.id),
                },
            )
        db.commit()
        return SourceBackedRuntimeSeed(
            organization_id=organization.id,
            actor_id=actor.id,
            bundle_id=bundle.id,
            storage=storage,
        )


def test_worker_claims_dispatches_and_completes(
    session_factory: sessionmaker[Session],
) -> None:
    seed = _seed_job(session_factory, slug="success")
    seen: list[JobExecutionContext] = []

    def handler(context: JobExecutionContext) -> dict[str, object]:
        seen.append(context)
        return {"processed": True, "source_file_id": context.payload["source_file_id"]}

    worker = ReliableJobWorker(
        session_factory=session_factory,
        handlers={"document_extraction": handler},
        worker_id="worker-success",
        lease_duration_seconds=30,
        heartbeat_interval_seconds=10,
    )
    assert worker.run_once() is True
    assert worker.run_once() is False
    assert len(seen) == 1

    with session_factory() as db:
        job = db.query(TaskJob).filter_by(id=seed.job_id).one()
        attempt = db.query(TaskJobAttempt).filter_by(job_id=seed.job_id).one()
        assert job.status == "completed"
        assert job.result_payload["processed"] is True
        assert attempt.status == "succeeded"
        assert attempt.worker_id == "worker-success"
        assert (
            db.query(AuditEvent)
            .filter_by(entity_id=seed.job_id, event_name="TaskJobCompleted")
            .count()
            == 1
        )


def test_worker_failure_is_recorded_and_dead_lettered(
    session_factory: sessionmaker[Session],
) -> None:
    seed = _seed_job(session_factory, slug="failure", max_attempts=1)

    def failing_handler(_context: JobExecutionContext) -> dict[str, object]:
        raise RuntimeError("deterministic parser failure")

    worker = ReliableJobWorker(
        session_factory=session_factory,
        handlers={"document_extraction": failing_handler},
        worker_id="worker-failure",
        lease_duration_seconds=30,
        heartbeat_interval_seconds=10,
    )
    assert worker.run_once() is True

    with session_factory() as db:
        job = db.query(TaskJob).filter_by(id=seed.job_id).one()
        attempt = db.query(TaskJobAttempt).filter_by(job_id=seed.job_id).one()
        assert job.status == "dead_letter"
        assert job.last_error_code == "job_handler_failed"
        assert attempt.status == "failed"
        assert attempt.error_code == "job_handler_failed"
        assert "deterministic parser failure" in attempt.error_message


def test_worker_does_not_retry_a_typed_permanent_failure(
    session_factory: sessionmaker[Session],
) -> None:
    seed = _seed_job(session_factory, slug="permanent-failure", max_attempts=3)

    def rejecting_handler(_context: JobExecutionContext) -> dict[str, object]:
        raise JobHandlerFailure("source_checksum_mismatch", "Checksum mismatch.")

    worker = ReliableJobWorker(
        session_factory=session_factory,
        handlers={"document_extraction": rejecting_handler},
        worker_id="worker-permanent-failure",
        lease_duration_seconds=30,
        heartbeat_interval_seconds=10,
    )
    assert worker.run_once() is True

    with session_factory() as db:
        job = db.query(TaskJob).filter_by(id=seed.job_id).one()
        attempt = db.query(TaskJobAttempt).filter_by(job_id=seed.job_id).one()
        assert job.status == "dead_letter"
        assert job.attempt_count == 1
        assert attempt.error_code == "source_checksum_mismatch"


def test_production_registry_extracts_real_sources_then_aligns(
    session_factory: sessionmaker[Session],
) -> None:
    seed = _seed_source_backed_jobs(session_factory, slug="source-backed")
    worker = ReliableJobWorker(
        session_factory=session_factory,
        handlers=build_handler_registry(
            session_factory=session_factory,
            storage=seed.storage,
        ),
        worker_id="worker-source-backed",
        lease_duration_seconds=30,
        heartbeat_interval_seconds=10,
    )
    assert worker.run_once() is True
    assert worker.run_once() is True

    with session_factory() as db:
        snapshots = (
            db.query(DossierExtractionSnapshot)
            .filter_by(
                organization_id=seed.organization_id,
                dossier_bundle_id=seed.bundle_id,
            )
            .all()
        )
        by_kind = {snapshot.source_kind: snapshot for snapshot in snapshots}
        assert set(by_kind) == {"excel", "docx"}
        actor = db.query(User).filter_by(id=seed.actor_id).one()
        alignment_job = enqueue_durable_job(
            db,
            actor=actor,
            org_id=seed.organization_id,
            job_type="dossier_alignment",
            idempotency_key="source-backed-alignment",
            payload={
                "dossier_bundle_id": str(seed.bundle_id),
                "excel_snapshot_id": str(by_kind["excel"].id),
                "report_snapshot_id": str(by_kind["docx"].id),
            },
        )
        db.commit()
        alignment_job_id = alignment_job.id

    assert worker.run_once() is True
    assert worker.run_once() is False

    with session_factory() as db:
        job = db.query(TaskJob).filter_by(id=alignment_job_id).one()
        run = db.query(DossierAlignmentRun).filter_by(dossier_bundle_id=seed.bundle_id).one()
        rows = db.query(DossierRowAlignment).filter_by(alignment_run_id=run.id).all()
        assert job.status == "completed"
        assert job.result_payload["alignment_run_id"] == str(run.id)
        assert run.total_excel_rows == 1
        assert len(rows) == 1
        assert rows[0].state in {"candidate", "review_required"}
        assert rows[0].reviewed_by_user_id is None
        assert rows[0].reviewed_at is None
        assert (
            db.query(AuditEvent)
            .filter_by(event_name="DossierSourceExtracted")
            .count()
            == 2
        )
        assert (
            db.query(AuditEvent)
            .filter_by(event_name="DossierRowAlignmentCandidatesGenerated")
            .count()
            == 1
        )
